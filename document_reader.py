import os
import re
import zipfile
from pathlib import Path
from tempfile import TemporaryDirectory

from modules.pdf_reader import leer_pdf, detectar_renglones


def limpiar(texto):
    texto = texto or ""
    texto = texto.replace("\x00", " ")
    texto = re.sub(r"\s+", " ", texto)
    return texto.strip()


def leer_txt(path):
    for enc in ["utf-8", "latin-1", "cp1252"]:
        try:
            return limpiar(Path(path).read_text(encoding=enc, errors="ignore"))
        except Exception:
            continue
    return ""


def leer_docx(path):
    try:
        import docx
        doc = docx.Document(path)
        partes = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                partes.append(" | ".join(cell.text for cell in row.cells))
        return limpiar(" ".join(partes))
    except Exception as e:
        return f"ERROR_LEYENDO_DOCX: {e}"


def leer_xlsx(path):
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        partes = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                vals = [str(v) for v in row if v is not None and str(v).strip()]
                if vals:
                    partes.append(" | ".join(vals))
        return limpiar(" ".join(partes))
    except Exception as e:
        return f"ERROR_LEYENDO_XLSX: {e}"


def leer_archivo(path):
    path = str(path)
    ext = Path(path).suffix.lower()

    if ext == ".pdf":
        return leer_pdf(path)
    if ext == ".docx":
        return leer_docx(path)
    if ext == ".xlsx":
        return leer_xlsx(path)
    if ext in [".txt", ".csv"]:
        return leer_txt(path)
    return ""


def extraer_documentos(path):
    path = str(path)
    ext = Path(path).suffix.lower()

    if ext == ".zip":
        docs = []
        try:
            with TemporaryDirectory() as tmp:
                with zipfile.ZipFile(path) as z:
                    z.extractall(tmp)
                for root, _, files in os.walk(tmp):
                    for name in files:
                        p = os.path.join(root, name)
                        texto = leer_archivo(p)
                        if texto:
                            docs.append({"archivo": name, "texto": texto, "renglones": detectar_renglones(texto)})
        except Exception as e:
            docs.append({"archivo": path, "texto": f"ERROR_LEYENDO_ZIP: {e}", "renglones": []})
        return docs

    texto = leer_archivo(path)
    return [{"archivo": path, "texto": texto, "renglones": detectar_renglones(texto)}] if texto else []


def extraer_renglones_archivos(archivos):
    salida = []
    for archivo in archivos or []:
        for doc in extraer_documentos(archivo):
            for renglon in doc.get("renglones", []):
                salida.append({"archivo": doc.get("archivo", archivo), "renglon": renglon})
    return salida
